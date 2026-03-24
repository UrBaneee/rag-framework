"""DOCX parser using python-docx — extracts paragraphs and table cells."""

import hashlib
import logging
from pathlib import Path

from rag.core.contracts.document import Document
from rag.core.contracts.ir_block import BlockType, IRBlock
from rag.core.contracts.parse_report import ParseReport
from rag.core.interfaces.parser import BaseParser

logger = logging.getLogger(__name__)

_PARSER_NAME = "docx"
_SUPPORTED_MIME_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)
_MIN_BLOCK_CHARS = 3

try:
    import docx as _docx_lib  # python-docx

    _DOCX_AVAILABLE = True
except ImportError:  # pragma: no cover
    _DOCX_AVAILABLE = False


class DocxParser(BaseParser):
    """Parser for .docx files using python-docx.

    Extracts body paragraphs and table cell text as IRBlocks. Headings are
    tagged with BlockType.HEADING; all other non-empty paragraphs become
    BlockType.PARAGRAPH. Table cells are concatenated per row and emitted as
    PARAGRAPH blocks.

    Raises:
        ImportError: If python-docx is not installed (raised at instantiation).

    Usage::

        parser = DocxParser()
        doc = parser.parse("/path/to/file.docx")
    """

    def __init__(self) -> None:
        if not _DOCX_AVAILABLE:
            raise ImportError(
                "python-docx is not installed. Install with: pip install python-docx"
            )

    def supports(self, mime_type: str) -> bool:
        return mime_type == _SUPPORTED_MIME_TYPE

    def parse(self, source_path: str) -> Document:
        """Parse a .docx file into a Document.

        Args:
            source_path: Absolute path to the .docx file.

        Returns:
            Document with IRBlocks extracted from paragraphs and tables.

        Raises:
            ValueError: If the file cannot be opened.
        """
        path = Path(source_path)
        try:
            doc = _docx_lib.Document(source_path)
        except Exception as exc:
            raise ValueError(f"Cannot open DOCX '{source_path}': {exc}") from exc

        blocks: list[IRBlock] = []

        # Extract body paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if len(text) < _MIN_BLOCK_CHARS:
                continue
            style_name = (para.style.name or "").lower()
            block_type = (
                BlockType.HEADING if "heading" in style_name else BlockType.PARAGRAPH
            )
            blocks.append(IRBlock(block_type=block_type, text=text))

        # Extract table cells
        for table in doc.tables:
            for row in table.rows:
                row_texts = [cell.text.strip() for cell in row.cells]
                row_text = " | ".join(t for t in row_texts if t)
                if len(row_text) >= _MIN_BLOCK_CHARS:
                    blocks.append(IRBlock(block_type=BlockType.PARAGRAPH, text=row_text))

        all_text = " ".join(b.text for b in blocks)
        report = ParseReport(
            char_count=len(all_text),
            block_count=len(blocks),
            non_printable_ratio=0.0,
            repetition_score=0.0,
            parser_used=_PARSER_NAME,
            fallback_triggered=False,
        )

        doc_id = hashlib.sha256(source_path.encode()).hexdigest()[:16]
        return Document(
            doc_id=doc_id,
            source_path=source_path,
            mime_type=_SUPPORTED_MIME_TYPE,
            metadata={"filename": path.name, "extension": "docx"},
            blocks=blocks,
            parse_report=report,
        )
